import streamlit as st
import google.generativeai as genai
import requests
import json
import time
import io
import os
import re
from datetime import datetime

# ── optional heavy deps ───────────────────────────────────────────────────────
try:
    import fitz  # PyMuPDF
    PYMUPDF_OK = True
except ImportError:
    PYMUPDF_OK = False

try:
    import textstat
    TEXTSTAT_OK = True
except ImportError:
    TEXTSTAT_OK = False

try:
    import language_tool_python
    LT_OK = True
except ImportError:
    LT_OK = False

try:
    from docx import Document as DocxDocument
    from docx.shared import Pt, Inches
    DOCX_OK = True
except ImportError:
    DOCX_OK = False

try:
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet
    REPORTLAB_OK = True
except ImportError:
    REPORTLAB_OK = False

try:
    from deep_translator import GoogleTranslator
    TRANSLATOR_OK = True
except ImportError:
    TRANSLATOR_OK = False

try:
    import plotly.graph_objects as go
    import plotly.express as px
    PLOTLY_OK = True
except ImportError:
    PLOTLY_OK = False

try:
    import pandas as pd
    PANDAS_OK = True
except ImportError:
    PANDAS_OK = False

# ── page config ───────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="ResearchPaper AI Studio",
    page_icon="🔬",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ── CSS ───────────────────────────────────────────────────────────────────────
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700&family=Merriweather:wght@400;700&display=swap');

:root {
    --primary: #2563EB;
    --primary-light: #DBEAFE;
    --accent: #7C3AED;
    --success: #059669;
    --warning: #D97706;
    --bg: #F8FAFC;
    --card: #FFFFFF;
    --border: #E2E8F0;
    --text: #1E293B;
    --muted: #64748B;
    --radius: 12px;
}

html, body, [class*="css"] {
    font-family: 'Inter', sans-serif;
    color: var(--text);
}

.stApp { background: var(--bg); }

.app-header {
    background: linear-gradient(135deg, #1E40AF 0%, #7C3AED 100%);
    padding: 1.5rem 2rem;
    border-radius: 0 0 20px 20px;
    margin-bottom: 1.5rem;
    display: flex;
    align-items: center;
    gap: 1rem;
}
.app-header h1 {
    color: white;
    font-family: 'Merriweather', serif;
    font-size: 1.6rem;
    font-weight: 700;
    margin: 0;
}
.app-header p { color: rgba(255,255,255,0.8); margin: 0; font-size: 0.85rem; }
.header-logo { font-size: 2.5rem; }

.card {
    background: var(--card);
    border: 1px solid var(--border);
    border-radius: var(--radius);
    padding: 1.25rem;
    margin-bottom: 1rem;
    box-shadow: 0 1px 3px rgba(0,0,0,0.06);
}
.card-title {
    font-weight: 600;
    font-size: 0.95rem;
    color: var(--primary);
    margin-bottom: 0.5rem;
    display: flex;
    align-items: center;
    gap: 0.4rem;
}

.metric-card {
    background: var(--card);
    border: 1px solid var(--border);
    border-radius: var(--radius);
    padding: 1rem;
    text-align: center;
    box-shadow: 0 1px 3px rgba(0,0,0,0.06);
}
.metric-value {
    font-size: 2rem;
    font-weight: 700;
    color: var(--primary);
    line-height: 1;
}
.metric-label { font-size: 0.78rem; color: var(--muted); margin-top: 0.3rem; }

.topic-card {
    background: var(--card);
    border: 1px solid var(--border);
    border-left: 4px solid var(--primary);
    border-radius: var(--radius);
    padding: 1rem 1.2rem;
    margin-bottom: 0.75rem;
    transition: box-shadow 0.2s;
}
.topic-card:hover { box-shadow: 0 4px 12px rgba(37,99,235,0.12); }
.topic-title { font-weight: 600; font-size: 0.95rem; margin-bottom: 0.4rem; }
.badge {
    display: inline-block;
    padding: 0.2rem 0.6rem;
    border-radius: 9999px;
    font-size: 0.72rem;
    font-weight: 500;
    margin-right: 0.3rem;
}
.badge-blue { background: #DBEAFE; color: #1D4ED8; }
.badge-purple { background: #EDE9FE; color: #6D28D9; }
.badge-green { background: #D1FAE5; color: #065F46; }
.badge-orange { background: #FEF3C7; color: #92400E; }

.ref-card {
    background: #F8FAFC;
    border: 1px solid var(--border);
    border-radius: 8px;
    padding: 0.9rem;
    margin-bottom: 0.6rem;
}
.ref-title { font-weight: 600; font-size: 0.9rem; }
.ref-meta { font-size: 0.78rem; color: var(--muted); margin-top: 0.2rem; }

.section-header {
    font-family: 'Merriweather', serif;
    font-size: 1.2rem;
    font-weight: 700;
    color: var(--text);
    border-bottom: 2px solid var(--primary-light);
    padding-bottom: 0.5rem;
    margin-bottom: 1rem;
}

.progress-step {
    display: flex;
    align-items: center;
    gap: 0.75rem;
    padding: 0.5rem 0;
}
.step-dot {
    width: 28px;
    height: 28px;
    border-radius: 50%;
    display: flex;
    align-items: center;
    justify-content: center;
    font-size: 0.75rem;
    font-weight: 700;
    flex-shrink: 0;
}
.step-done { background: #D1FAE5; color: #065F46; }
.step-active { background: var(--primary); color: white; }
.step-todo { background: var(--border); color: var(--muted); }
.step-label { font-size: 0.85rem; }

.tip-box {
    background: #EFF6FF;
    border-left: 3px solid var(--primary);
    border-radius: 0 8px 8px 0;
    padding: 0.8rem 1rem;
    font-size: 0.83rem;
    color: #1E40AF;
    margin-bottom: 1rem;
}

.warning-box {
    background: #FFFBEB;
    border-left: 3px solid var(--warning);
    border-radius: 0 8px 8px 0;
    padding: 0.8rem 1rem;
    font-size: 0.83rem;
    color: #92400E;
    margin-bottom: 1rem;
}

div[data-testid="stSidebarNav"] { display: none; }
.stButton > button {
    border-radius: 8px;
    font-weight: 500;
    transition: all 0.2s;
}
.stButton > button:hover { transform: translateY(-1px); box-shadow: 0 4px 12px rgba(37,99,235,0.2); }
</style>
""", unsafe_allow_html=True)

# ── session state init ────────────────────────────────────────────────────────
DEFAULTS = {
    "page": "Dashboard",
    "selected_topic": None,
    "topics": [],
    "references": [],
    "pdf_analyses": [],
    "research_gaps": None,
    "methodology": None,
    "paper_sections": {},
    "literature_matrix": [],
    "api_key_set": False,
}
for k, v in DEFAULTS.items():
    if k not in st.session_state:
        st.session_state[k] = v

# ── Gemini client ─────────────────────────────────────────────────────────────
SYSTEM_INSTRUCTION = (
    "You are an academic research assistant. "
    "Use ONLY the information provided. "
    "Do not invent references, authors, datasets, experiments, statistics, or results. "
    "If information is unavailable, state: 'Information not provided.'"
)

def get_gemini_client():
    api_key = os.environ.get("GEMINI_API_KEY") or st.session_state.get("gemini_api_key", "")
    if not api_key:
        return None
    genai.configure(api_key=api_key)
    return genai.GenerativeModel(
        model_name="gemini-2.0-flash-lite",
        system_instruction=SYSTEM_INSTRUCTION,
    )

def call_gemini(prompt: str, json_mode: bool = False) -> str:
    model = get_gemini_client()
    if model is None:
        st.error("⚠️ Gemini API key not set. Please enter it in the sidebar.")
        return ""
    config = {"temperature": 0.4, "max_output_tokens": 2048}
    if json_mode:
        config["response_mime_type"] = "application/json"
    for attempt in range(3):
        try:
            response = model.generate_content(prompt, generation_config=config)
            return response.text
        except Exception as e:
            if attempt == 2:
                st.error(f"Gemini error: {e}")
                return ""
            time.sleep(1.5 * (attempt + 1))
    return ""

def parse_json_response(text: str):
    try:
        clean = re.sub(r"```(?:json)?", "", text).strip().rstrip("`")
        return json.loads(clean)
    except Exception:
        return None

# ── research APIs ─────────────────────────────────────────────────────────────
def search_arxiv(query: str, max_results: int = 5):
    url = "http://export.arxiv.org/api/query"
    params = {"search_query": f"all:{query}", "max_results": max_results, "sortBy": "relevance"}
    try:
        r = requests.get(url, params=params, timeout=10)
        r.raise_for_status()
        entries = []
        import xml.etree.ElementTree as ET
        root = ET.fromstring(r.text)
        ns = {"a": "http://www.w3.org/2005/Atom"}
        for entry in root.findall("a:entry", ns):
            title = entry.find("a:title", ns).text.strip().replace("\n", " ")
            authors = [a.find("a:name", ns).text for a in entry.findall("a:author", ns)]
            year_str = entry.find("a:published", ns).text[:4]
            link = entry.find("a:id", ns).text.strip()
            abstract = entry.find("a:summary", ns).text.strip().replace("\n", " ")
            entries.append({
                "title": title,
                "authors": authors,
                "year": year_str,
                "doi": link,
                "abstract": abstract[:300] + "…" if len(abstract) > 300 else abstract,
                "source": "ArXiv",
            })
        return entries
    except Exception as e:
        st.warning(f"ArXiv search failed: {e}")
        return []

def search_crossref(query: str, max_results: int = 5):
    url = "https://api.crossref.org/works"
    params = {"query": query, "rows": max_results, "select": "title,author,published,DOI,abstract"}
    try:
        r = requests.get(url, params=params, timeout=10)
        r.raise_for_status()
        items = r.json().get("message", {}).get("items", [])
        entries = []
        for item in items:
            title = " ".join(item.get("title", ["Untitled"]))
            authors = [f"{a.get('given','')} {a.get('family','')}".strip() for a in item.get("author", [])]
            year = str(item.get("published", {}).get("date-parts", [[""]])[0][0])
            doi = item.get("DOI", "")
            abstract = item.get("abstract", "")
            if abstract:
                abstract = re.sub(r"<[^>]+>", "", abstract)[:300] + "…"
            entries.append({
                "title": title,
                "authors": authors,
                "year": year,
                "doi": doi,
                "abstract": abstract,
                "source": "CrossRef",
            })
        return entries
    except Exception as e:
        st.warning(f"CrossRef search failed: {e}")
        return []

# ── PDF extraction ────────────────────────────────────────────────────────────
def extract_pdf_text(uploaded_file) -> str:
    if not PYMUPDF_OK:
        return ""
    data = uploaded_file.read()
    doc = fitz.open(stream=data, filetype="pdf")
    text = ""
    for page in doc:
        text += page.get_text()
    return text[:8000]

# ── completion tracker ────────────────────────────────────────────────────────
def completion_pct() -> int:
    steps = [
        bool(st.session_state.selected_topic),
        len(st.session_state.references) > 0,
        len(st.session_state.pdf_analyses) > 0,
        bool(st.session_state.research_gaps),
        bool(st.session_state.methodology),
        len(st.session_state.paper_sections) >= 3,
        len(st.session_state.paper_sections) >= 7,
    ]
    return int(sum(steps) / len(steps) * 100)

# ── sidebar ───────────────────────────────────────────────────────────────────
NAV_ITEMS = [
    ("🏠", "Dashboard"),
    ("💡", "Topic Generator"),
    ("📚", "Reference Finder"),
    ("📄", "PDF Analyzer"),
    ("📊", "Literature Matrix"),
    ("🔍", "Research Gap Finder"),
    ("🧪", "Methodology Builder"),
    ("✍️", "Paper Builder"),
    ("🔖", "Citation Manager"),
    ("📝", "Writing Assistant"),
    ("🌍", "Translation"),
    ("📤", "Export"),
]

with st.sidebar:
    st.markdown("### 🔬 ResearchPaper AI")
    st.divider()

    api_key = st.text_input("Gemini API Key", type="password",
                             value=st.session_state.get("gemini_api_key", ""),
                             placeholder="Enter your key…")
    if api_key:
        st.session_state["gemini_api_key"] = api_key
        st.session_state["api_key_set"] = True

    st.divider()
    st.markdown("**Navigation**")
    for icon, label in NAV_ITEMS:
        if st.button(f"{icon} {label}", key=f"nav_{label}", use_container_width=True):
            st.session_state.page = label

    st.divider()
    pct = completion_pct()
    st.markdown(f"**Overall Progress** — {pct}%")
    st.progress(pct / 100)

    steps = [
        ("Topic Selected", bool(st.session_state.selected_topic)),
        ("References Added", len(st.session_state.references) > 0),
        ("PDFs Analyzed", len(st.session_state.pdf_analyses) > 0),
        ("Gaps Found", bool(st.session_state.research_gaps)),
        ("Methodology Set", bool(st.session_state.methodology)),
        ("Sections Written", len(st.session_state.paper_sections) >= 3),
    ]
    for label, done in steps:
        icon = "✅" if done else "⬜"
        st.markdown(f"{icon} {label}")

# ── header ────────────────────────────────────────────────────────────────────
st.markdown("""
<div class="app-header">
  <div class="header-logo">🔬</div>
  <div>
    <h1>ResearchPaper AI Studio</h1>
    <p>Your intelligent companion for writing your first research paper — step by step.</p>
  </div>
</div>
""", unsafe_allow_html=True)

page = st.session_state.page

# ═══════════════════════════════════════════════════════════════════════════════
# PAGE: DASHBOARD
# ═══════════════════════════════════════════════════════════════════════════════
if page == "Dashboard":
    topic = st.session_state.selected_topic
    topic_title = topic["title"] if topic else "Not selected yet"

    c1, c2, c3, c4 = st.columns(4)
    with c1:
        st.markdown(f"""<div class="metric-card">
            <div class="metric-value">{len(st.session_state.references)}</div>
            <div class="metric-label">References Saved</div>
        </div>""", unsafe_allow_html=True)
    with c2:
        st.markdown(f"""<div class="metric-card">
            <div class="metric-value">{len(st.session_state.pdf_analyses)}</div>
            <div class="metric-label">PDFs Analyzed</div>
        </div>""", unsafe_allow_html=True)
    with c3:
        st.markdown(f"""<div class="metric-card">
            <div class="metric-value">{len(st.session_state.paper_sections)}</div>
            <div class="metric-label">Sections Written</div>
        </div>""", unsafe_allow_html=True)
    with c4:
        st.markdown(f"""<div class="metric-card">
            <div class="metric-value">{completion_pct()}%</div>
            <div class="metric-label">Complete</div>
        </div>""", unsafe_allow_html=True)

    st.markdown("<br>", unsafe_allow_html=True)
    col_left, col_right = st.columns([1.4, 1])

    with col_left:
        st.markdown('<div class="section-header">📋 Selected Topic</div>', unsafe_allow_html=True)
        if topic:
            st.markdown(f"""<div class="card">
                <div class="card-title">🎯 {topic['title']}</div>
                <p style="font-size:0.85rem;color:#475569;margin-top:0.4rem">{topic.get('problem_statement','')}</p>
                <span class="badge badge-blue">{topic.get('methodology_type','')}</span>
                <span class="badge badge-purple">Difficulty: {topic.get('difficulty_score','')}/10</span>
                <span class="badge badge-green">~{topic.get('expected_pages','')} pages</span>
            </div>""", unsafe_allow_html=True)
        else:
            st.markdown('<div class="tip-box">💡 Start by generating a research topic in the <b>Topic Generator</b>.</div>', unsafe_allow_html=True)

        st.markdown('<div class="section-header">🗺️ Research Pipeline</div>', unsafe_allow_html=True)
        pipeline = [
            ("Topic Selection", bool(st.session_state.selected_topic)),
            ("Reference Finder", len(st.session_state.references) > 0),
            ("PDF Analysis", len(st.session_state.pdf_analyses) > 0),
            ("Literature Matrix", len(st.session_state.literature_matrix) > 0),
            ("Research Gap Finder", bool(st.session_state.research_gaps)),
            ("Methodology Builder", bool(st.session_state.methodology)),
            ("Paper Builder", len(st.session_state.paper_sections) > 0),
            ("Citation Manager", False),
            ("Export", False),
        ]
        for i, (label, done) in enumerate(pipeline):
            css = "step-done" if done else ("step-active" if i == next((j for j, (_, d) in enumerate(pipeline) if not d), len(pipeline)) else "step-todo")
            icon = "✓" if done else str(i + 1)
            st.markdown(f"""<div class="progress-step">
                <div class="step-dot {css}">{icon}</div>
                <div class="step-label">{label}</div>
            </div>""", unsafe_allow_html=True)

    with col_right:
        st.markdown('<div class="section-header">📚 Saved References</div>', unsafe_allow_html=True)
        if st.session_state.references:
            for ref in st.session_state.references[-5:]:
                st.markdown(f"""<div class="ref-card">
                    <div class="ref-title">{ref['title'][:60]}…</div>
                    <div class="ref-meta">{', '.join(ref['authors'][:2])} · {ref['year']} · {ref['source']}</div>
                </div>""", unsafe_allow_html=True)
        else:
            st.markdown('<div class="tip-box">No references yet. Use the Reference Finder to search ArXiv and CrossRef.</div>', unsafe_allow_html=True)

        if PLOTLY_OK and len(st.session_state.paper_sections) > 0:
            st.markdown('<div class="section-header">📄 Sections Progress</div>', unsafe_allow_html=True)
            all_sections = ["Abstract", "Introduction", "Literature Review", "Methodology", "Results", "Discussion", "Conclusion", "Future Work"]
            done_secs = [s for s in all_sections if s in st.session_state.paper_sections]
            not_done = [s for s in all_sections if s not in st.session_state.paper_sections]
            fig = go.Figure(go.Bar(
                x=[len(done_secs), len(not_done)],
                y=["Written", "Remaining"],
                orientation="h",
                marker_color=["#2563EB", "#E2E8F0"],
            ))
            fig.update_layout(height=120, margin=dict(l=0, r=0, t=0, b=0), paper_bgcolor="rgba(0,0,0,0)", plot_bgcolor="rgba(0,0,0,0)", showlegend=False)
            st.plotly_chart(fig, use_container_width=True)

# ═══════════════════════════════════════════════════════════════════════════════
# PAGE: TOPIC GENERATOR
# ═══════════════════════════════════════════════════════════════════════════════
elif page == "Topic Generator":
    st.markdown('<div class="section-header">💡 Topic Generator</div>', unsafe_allow_html=True)
    st.markdown('<div class="tip-box">Enter your field and keywords, and the AI will suggest 5–10 original research topics with difficulty and novelty scores.</div>', unsafe_allow_html=True)

    c1, c2 = st.columns(2)
    with c1:
        field = st.text_input("Research Field *", placeholder="e.g. Computer Science, Biology…")
        domain = st.text_input("Domain / Sub-field", placeholder="e.g. Machine Learning, Genomics…")
    with c2:
        difficulty = st.selectbox("Your Level", ["Beginner (UG Year 1-2)", "Intermediate (UG Year 3-4)", "Advanced (PG)"])
        keywords = st.text_input("Keywords (comma-separated)", placeholder="e.g. NLP, sentiment, social media")

    if st.button("🚀 Generate Topics", type="primary", disabled=not field):
        prompt = f"""Generate 7 original research topic ideas for a student.
Field: {field}
Domain: {domain or 'General'}
Level: {difficulty}
Keywords: {keywords or 'not specified'}

Respond ONLY with a JSON array. Each element has:
- title (string)
- problem_statement (string, 2 sentences)
- research_gap (string, 1 sentence)
- novelty_score (integer 1-10)
- difficulty_score (integer 1-10)
- methodology_type (string: Survey|ML|Experimental|Comparative|Case Study)
- expected_pages (integer)

No extra text. Pure JSON."""
        with st.spinner("Generating topics…"):
            raw = call_gemini(prompt, json_mode=True)
        data = parse_json_response(raw)
        if data:
            st.session_state.topics = data if isinstance(data, list) else data.get("topics", [])
        else:
            st.error("Could not parse topics. Please try again.")

    if st.session_state.topics:
        st.markdown(f"### {len(st.session_state.topics)} Topics Generated")
        for i, t in enumerate(st.session_state.topics):
            with st.container():
                st.markdown(f"""<div class="topic-card">
                    <div class="topic-title">{i+1}. {t.get('title','')}</div>
                    <p style="font-size:0.84rem;color:#475569;margin:0.3rem 0">{t.get('problem_statement','')}</p>
                    <p style="font-size:0.82rem;color:#64748B;font-style:italic">Gap: {t.get('research_gap','')}</p>
                    <span class="badge badge-blue">{t.get('methodology_type','')}</span>
                    <span class="badge badge-purple">Novelty {t.get('novelty_score','')}/10</span>
                    <span class="badge badge-orange">Difficulty {t.get('difficulty_score','')}/10</span>
                    <span class="badge badge-green">~{t.get('expected_pages','')} pages</span>
                </div>""", unsafe_allow_html=True)
                if st.button(f"✅ Select This Topic", key=f"sel_{i}"):
                    st.session_state.selected_topic = t
                    st.success(f"Topic selected: {t['title']}")
                    st.rerun()

# ═══════════════════════════════════════════════════════════════════════════════
# PAGE: REFERENCE FINDER
# ═══════════════════════════════════════════════════════════════════════════════
elif page == "Reference Finder":
    st.markdown('<div class="section-header">📚 Reference Finder</div>', unsafe_allow_html=True)
    st.markdown('<div class="tip-box">Search real academic databases (ArXiv + CrossRef) to find genuine references for your topic.</div>', unsafe_allow_html=True)

    default_q = st.session_state.selected_topic["title"] if st.session_state.selected_topic else ""
    query = st.text_input("Search query", value=default_q, placeholder="Enter keywords to search…")
    c1, c2, c3 = st.columns(3)
    with c1:
        max_res = st.slider("Max results per source", 3, 10, 5)
    with c2:
        use_arxiv = st.checkbox("Search ArXiv", value=True)
    with c3:
        use_crossref = st.checkbox("Search CrossRef", value=True)

    if st.button("🔍 Search", type="primary", disabled=not query):
        results = []
        with st.spinner("Searching academic databases…"):
            if use_arxiv:
                results += search_arxiv(query, max_res)
            if use_crossref:
                results += search_crossref(query, max_res)
        st.session_state["search_results"] = results

    results = st.session_state.get("search_results", [])
    if results:
        st.markdown(f"**{len(results)} papers found**")
        saved_dois = {r["doi"] for r in st.session_state.references}
        for r in results:
            already = r["doi"] in saved_dois
            with st.expander(f"📄 {r['title'][:80]}", expanded=False):
                st.markdown(f"""<div class="ref-card">
                    <div class="ref-title">{r['title']}</div>
                    <div class="ref-meta">👤 {', '.join(r['authors'][:3])} · 📅 {r['year']} · 🔗 {r['source']}</div>
                    <p style="font-size:0.82rem;margin-top:0.5rem;color:#475569">{r['abstract']}</p>
                    {f'<div style="font-size:0.75rem;color:#2563EB">DOI: {r["doi"]}</div>' if r["doi"] else ''}
                </div>""", unsafe_allow_html=True)
                if already:
                    st.success("✓ Already in your library")
                elif st.button("➕ Add to Library", key=f"add_{r['doi']}_{r['title'][:20]}"):
                    st.session_state.references.append(r)
                    st.success("Added!")
                    st.rerun()

    if st.session_state.references:
        st.divider()
        st.markdown(f"### 📖 Your Reference Library ({len(st.session_state.references)} papers)")
        for i, ref in enumerate(st.session_state.references):
            col1, col2 = st.columns([10, 1])
            with col1:
                st.markdown(f"""<div class="ref-card">
                    <div class="ref-title">{ref['title']}</div>
                    <div class="ref-meta">{', '.join(ref['authors'][:3])} · {ref['year']} · {ref['source']}</div>
                </div>""", unsafe_allow_html=True)
            with col2:
                if st.button("🗑️", key=f"del_{i}"):
                    st.session_state.references.pop(i)
                    st.rerun()

# ═══════════════════════════════════════════════════════════════════════════════
# PAGE: PDF ANALYZER
# ═══════════════════════════════════════════════════════════════════════════════
elif page == "PDF Analyzer":
    st.markdown('<div class="section-header">📄 PDF Analyzer</div>', unsafe_allow_html=True)
    if not PYMUPDF_OK:
        st.warning("PyMuPDF not installed. Install with: `pip install pymupdf`")

    files = st.file_uploader("Upload research papers (PDF)", type=["pdf"], accept_multiple_files=True)

    if files and PYMUPDF_OK:
        if st.button("🔬 Analyze Papers", type="primary"):
            for f in files:
                if any(a["filename"] == f.name for a in st.session_state.pdf_analyses):
                    continue
                with st.spinner(f"Analyzing {f.name}…"):
                    text = extract_pdf_text(f)
                    if not text.strip():
                        st.warning(f"Could not extract text from {f.name}")
                        continue
                    prompt = f"""Analyze this academic paper extract and return ONLY JSON with these keys:
- summary (string, 3 sentences)
- objectives (list of strings)
- methodology (string)
- findings (string)
- limitations (string)
- future_work (string)
- authors_noted (string, mention only if clearly stated in text)
- year_noted (string, mention only if clearly stated)

Paper text:
{text[:4000]}

Return pure JSON only."""
                    raw = call_gemini(prompt, json_mode=True)
                    data = parse_json_response(raw)
                    if data:
                        data["filename"] = f.name
                        st.session_state.pdf_analyses.append(data)
                        # auto-add to literature matrix
                        st.session_state.literature_matrix.append({
                            "Paper": f.name,
                            "Authors": data.get("authors_noted", "Not stated"),
                            "Year": data.get("year_noted", "N/A"),
                            "Method": data.get("methodology", "N/A")[:80],
                            "Findings": data.get("findings", "N/A")[:100],
                            "Limitations": data.get("limitations", "N/A")[:80],
                        })
            st.success("Analysis complete!")

    if st.session_state.pdf_analyses:
        for analysis in st.session_state.pdf_analyses:
            with st.expander(f"📄 {analysis['filename']}", expanded=False):
                tabs = st.tabs(["Summary", "Objectives", "Methodology", "Findings", "Limitations", "Future Work"])
                with tabs[0]: st.write(analysis.get("summary", "N/A"))
                with tabs[1]:
                    for obj in analysis.get("objectives", []):
                        st.markdown(f"• {obj}")
                with tabs[2]: st.write(analysis.get("methodology", "N/A"))
                with tabs[3]: st.write(analysis.get("findings", "N/A"))
                with tabs[4]: st.write(analysis.get("limitations", "N/A"))
                with tabs[5]: st.write(analysis.get("future_work", "N/A"))

# ═══════════════════════════════════════════════════════════════════════════════
# PAGE: LITERATURE MATRIX
# ═══════════════════════════════════════════════════════════════════════════════
elif page == "Literature Matrix":
    st.markdown('<div class="section-header">📊 Literature Matrix</div>', unsafe_allow_html=True)
    st.markdown('<div class="tip-box">A structured comparison table of all analyzed papers. Analyze PDFs first to populate this table.</div>', unsafe_allow_html=True)

    matrix = st.session_state.literature_matrix
    if matrix and PANDAS_OK:
        df = pd.DataFrame(matrix)
        st.dataframe(df, use_container_width=True, height=300)
        csv = df.to_csv(index=False).encode("utf-8")
        st.download_button("📥 Download CSV", csv, "literature_matrix.csv", "text/csv")
    else:
        st.info("No papers analyzed yet. Upload PDFs in the PDF Analyzer section.")

    if st.session_state.references and PANDAS_OK:
        st.divider()
        st.markdown("### References Summary Table")
        ref_data = [{
            "Title": r["title"][:60] + "…" if len(r["title"]) > 60 else r["title"],
            "Authors": ", ".join(r["authors"][:2]),
            "Year": r["year"],
            "Source": r["source"],
            "DOI": r["doi"][:30] if r["doi"] else "N/A",
        } for r in st.session_state.references]
        df2 = pd.DataFrame(ref_data)
        st.dataframe(df2, use_container_width=True)

# ═══════════════════════════════════════════════════════════════════════════════
# PAGE: RESEARCH GAP FINDER
# ═══════════════════════════════════════════════════════════════════════════════
elif page == "Research Gap Finder":
    st.markdown('<div class="section-header">🔍 Research Gap Finder</div>', unsafe_allow_html=True)
    st.markdown('<div class="warning-box">⚠️ This analysis is based <b>only</b> on references and PDFs you have provided. The AI will not invent any information.</div>', unsafe_allow_html=True)

    refs_summary = "\n".join([f"- {r['title']} ({r['year']}) — {r['abstract'][:150]}" for r in st.session_state.references[:10]])
    pdf_summary = "\n".join([f"- {a['filename']}: {a.get('summary','')[:150]}" for a in st.session_state.pdf_analyses[:5]])

    if not refs_summary and not pdf_summary:
        st.info("Add references and analyze PDFs first to identify research gaps.")
    else:
        if st.button("🔍 Identify Research Gaps", type="primary"):
            prompt = f"""Based ONLY on the following references and paper analyses provided by the student, identify research gaps.
Do NOT invent any studies, statistics, or authors not mentioned below.

REFERENCES:
{refs_summary or 'None provided'}

PDF ANALYSES:
{pdf_summary or 'None provided'}

Return ONLY JSON with:
- existing_trends (list of strings, max 4)
- research_gaps (list of strings, max 5)
- opportunities (list of strings, max 4)
- recommended_direction (string, 2-3 sentences)

Pure JSON only."""
            with st.spinner("Analyzing gaps…"):
                raw = call_gemini(prompt, json_mode=True)
            data = parse_json_response(raw)
            if data:
                st.session_state.research_gaps = data
            else:
                st.error("Could not parse response. Try again.")

    if st.session_state.research_gaps:
        g = st.session_state.research_gaps
        c1, c2 = st.columns(2)
        with c1:
            st.markdown("#### 📈 Existing Trends")
            for t in g.get("existing_trends", []):
                st.markdown(f"<div class='card'>🔵 {t}</div>", unsafe_allow_html=True)
        with c2:
            st.markdown("#### 🕳️ Research Gaps")
            for gap in g.get("research_gaps", []):
                st.markdown(f"<div class='card'>🔴 {gap}</div>", unsafe_allow_html=True)

        st.markdown("#### 💡 Opportunities")
        for opp in g.get("opportunities", []):
            st.markdown(f"<div class='card'>🟢 {opp}</div>", unsafe_allow_html=True)

        st.markdown("#### 🎯 Recommended Direction")
        st.info(g.get("recommended_direction", ""))

# ═══════════════════════════════════════════════════════════════════════════════
# PAGE: METHODOLOGY BUILDER
# ═══════════════════════════════════════════════════════════════════════════════
elif page == "Methodology Builder":
    st.markdown('<div class="section-header">🧪 Methodology Builder</div>', unsafe_allow_html=True)
    st.markdown('<div class="tip-box">Select your research type and the AI will generate a complete methodology framework based on your topic.</div>', unsafe_allow_html=True)

    topic_title = st.session_state.selected_topic["title"] if st.session_state.selected_topic else ""
    c1, c2 = st.columns(2)
    with c1:
        topic_input = st.text_input("Research Topic", value=topic_title)
        research_type = st.selectbox("Research Type", ["Survey", "Machine Learning", "Experimental", "Comparative", "Case Study"])
    with c2:
        domain_input = st.text_input("Domain", placeholder="e.g. Healthcare, Education…")
        constraints = st.text_area("Constraints (optional)", placeholder="e.g. No primary data collection, only secondary sources…", height=80)

    if st.button("⚙️ Generate Methodology", type="primary", disabled=not topic_input):
        prompt = f"""Design a research methodology for a student's paper.
Topic: {topic_input}
Research Type: {research_type}
Domain: {domain_input or 'Not specified'}
Constraints: {constraints or 'None'}

Return ONLY JSON with:
- objectives (list of 3-4 strings)
- variables (object with "independent" list and "dependent" list)
- data_collection (list of strings)
- tools (list of strings)
- evaluation_metrics (list of strings)
- workflow (list of strings, ordered steps)
- limitations (list of strings)

Do not invent datasets or experiments. If unsure, describe general approaches only.
Pure JSON."""
        with st.spinner("Building methodology…"):
            raw = call_gemini(prompt, json_mode=True)
        data = parse_json_response(raw)
        if data:
            st.session_state.methodology = data
        else:
            st.error("Could not parse. Please try again.")

    if st.session_state.methodology:
        m = st.session_state.methodology
        tab1, tab2, tab3, tab4 = st.tabs(["Objectives & Variables", "Data & Tools", "Evaluation", "Workflow"])
        with tab1:
            st.markdown("**Objectives**")
            for o in m.get("objectives", []): st.markdown(f"• {o}")
            st.markdown("**Variables**")
            vars_ = m.get("variables", {})
            c1, c2 = st.columns(2)
            with c1:
                st.markdown("*Independent:*")
                for v in vars_.get("independent", []): st.markdown(f"→ {v}")
            with c2:
                st.markdown("*Dependent:*")
                for v in vars_.get("dependent", []): st.markdown(f"→ {v}")
        with tab2:
            c1, c2 = st.columns(2)
            with c1:
                st.markdown("**Data Collection**")
                for d in m.get("data_collection", []): st.markdown(f"• {d}")
            with c2:
                st.markdown("**Tools**")
                for t in m.get("tools", []): st.markdown(f"• {t}")
        with tab3:
            st.markdown("**Evaluation Metrics**")
            for e in m.get("evaluation_metrics", []): st.markdown(f"📊 {e}")
            st.markdown("**Limitations**")
            for l in m.get("limitations", []): st.markdown(f"⚠️ {l}")
        with tab4:
            st.markdown("**Research Workflow**")
            for i, step in enumerate(m.get("workflow", []), 1):
                st.markdown(f"""<div class="progress-step">
                    <div class="step-dot step-active">{i}</div>
                    <div class="step-label">{step}</div>
                </div>""", unsafe_allow_html=True)

# ═══════════════════════════════════════════════════════════════════════════════
# PAGE: PAPER BUILDER
# ═══════════════════════════════════════════════════════════════════════════════
elif page == "Paper Builder":
    st.markdown('<div class="section-header">✍️ Paper Builder</div>', unsafe_allow_html=True)
    st.markdown('<div class="warning-box">⚠️ The AI will NOT generate fake results, statistics, or experiments. Results sections will contain academic templates only.</div>', unsafe_allow_html=True)

    topic = st.session_state.selected_topic
    refs_text = "\n".join([f"- {r['title']} ({r['year']}) by {', '.join(r['authors'][:2])}" for r in st.session_state.references[:8]])
    meth_text = json.dumps(st.session_state.methodology, indent=2) if st.session_state.methodology else "Not generated yet."
    gaps_text = json.dumps(st.session_state.research_gaps, indent=2) if st.session_state.research_gaps else "Not generated yet."

    SECTIONS = [
        ("Abstract", "Generate a 250-word academic abstract"),
        ("Introduction", "Generate an introduction section"),
        ("Literature Review", "Generate a literature review based only on provided references"),
        ("Methodology", "Generate a methodology section based only on the provided methodology"),
        ("Results Framework", "Generate a results section TEMPLATE ONLY — no invented data"),
        ("Discussion", "Generate a discussion section based only on provided information"),
        ("Conclusion", "Generate a conclusion section"),
        ("Future Work", "Generate a future work section"),
    ]

    for sec_name, sec_desc in SECTIONS:
        with st.expander(f"{'✅' if sec_name in st.session_state.paper_sections else '⬜'} {sec_name}", expanded=False):
            col1, col2 = st.columns([3, 1])
            with col2:
                if st.button(f"Generate", key=f"gen_{sec_name}", type="primary"):
                    prompt = f"""{sec_desc} for a research paper.

Topic: {topic['title'] if topic else 'Not selected'}
Problem: {topic.get('problem_statement','') if topic else ''}
References provided:
{refs_text or 'None'}
Methodology provided:
{meth_text}
Research gaps provided:
{gaps_text}

IMPORTANT: Use ONLY the above information. Do not invent citations, data, or results.
For Results sections: write a TEMPLATE showing what format results should take, with placeholder notes.
Write formal academic prose. Length appropriate for the section."""
                    with st.spinner(f"Writing {sec_name}…"):
                        text = call_gemini(prompt)
                    if text:
                        st.session_state.paper_sections[sec_name] = text
            if sec_name in st.session_state.paper_sections:
                edited = st.text_area(
                    f"Edit {sec_name}",
                    value=st.session_state.paper_sections[sec_name],
                    height=250,
                    key=f"edit_{sec_name}",
                )
                st.session_state.paper_sections[sec_name] = edited
            else:
                st.info(f"Click 'Generate' to create the {sec_name}.")

# ═══════════════════════════════════════════════════════════════════════════════
# PAGE: CITATION MANAGER
# ═══════════════════════════════════════════════════════════════════════════════
elif page == "Citation Manager":
    st.markdown('<div class="section-header">🔖 Citation Manager</div>', unsafe_allow_html=True)
    st.markdown('<div class="warning-box">⚠️ Citations are generated ONLY from references in your library. No citations are invented.</div>', unsafe_allow_html=True)

    if not st.session_state.references:
        st.info("Add references in the Reference Finder first.")
    else:
        style = st.selectbox("Citation Style", ["APA", "MLA", "IEEE", "Chicago"])
        if st.button("🔖 Generate Citations", type="primary"):
            refs_data = json.dumps([{
                "title": r["title"],
                "authors": r["authors"],
                "year": r["year"],
                "doi": r["doi"],
                "source": r["source"],
            } for r in st.session_state.references], indent=2)
            prompt = f"""Format these references in {style} citation style.
Return ONLY JSON with:
- in_text (list of strings, one per reference, e.g. "(Smith, 2021)")
- full_citations (list of strings, one per reference, properly formatted)

References:
{refs_data}

Use ONLY the data provided. Do not invent anything. Pure JSON."""
            with st.spinner("Generating citations…"):
                raw = call_gemini(prompt, json_mode=True)
            data = parse_json_response(raw)
            if data:
                st.session_state["citations"] = data

        if st.session_state.get("citations"):
            c = st.session_state["citations"]
            tab1, tab2 = st.tabs(["In-Text Citations", "Reference List"])
            with tab1:
                for i, ref in enumerate(st.session_state.references):
                    cite = c["in_text"][i] if i < len(c.get("in_text", [])) else "N/A"
                    st.markdown(f"""<div class="ref-card">
                        <div class="ref-title">{ref['title'][:60]}</div>
                        <div style="font-family:monospace;background:#F1F5F9;padding:0.4rem;border-radius:6px;margin-top:0.4rem">{cite}</div>
                    </div>""", unsafe_allow_html=True)
            with tab2:
                full_text = "\n\n".join(c.get("full_citations", []))
                st.text_area("Reference List", value=full_text, height=300)
                st.download_button("📥 Copy Reference List", full_text, f"references_{style}.txt")

# ═══════════════════════════════════════════════════════════════════════════════
# PAGE: WRITING ASSISTANT
# ═══════════════════════════════════════════════════════════════════════════════
elif page == "Writing Assistant":
    st.markdown('<div class="section-header">📝 Writing Assistant</div>', unsafe_allow_html=True)

    text_input = st.text_area("Paste your draft text here for analysis:", height=200, placeholder="Paste your paragraph or section here…")

    if text_input.strip():
        col1, col2, col3 = st.columns(3)
        if TEXTSTAT_OK:
            with col1:
                fk = textstat.flesch_kincaid_grade(text_input)
                st.metric("Flesch-Kincaid Grade", f"{fk:.1f}")
            with col2:
                ease = textstat.flesch_reading_ease(text_input)
                st.metric("Reading Ease", f"{ease:.0f}/100")
            with col3:
                words = textstat.lexicon_count(text_input)
                st.metric("Word Count", words)

        if st.button("🔍 Analyze Writing", type="primary"):
            prompt = f"""Analyze this academic text draft and provide writing feedback.

Text:
{text_input[:2000]}

Return ONLY JSON with:
- academic_tone_score (integer 1-10)
- clarity_score (integer 1-10)
- suggestions (list of 3-5 strings, specific actionable improvements)
- passive_voice_issues (list of strings, specific phrases to change, or empty list)
- grammar_notes (list of strings)
- improved_sentence (string, one rewritten example sentence from the text)

Pure JSON."""
            with st.spinner("Analyzing…"):
                raw = call_gemini(prompt, json_mode=True)
            data = parse_json_response(raw)
            if data:
                st.session_state["writing_analysis"] = data

        if st.session_state.get("writing_analysis"):
            a = st.session_state["writing_analysis"]
            c1, c2 = st.columns(2)
            with c1:
                st.metric("Academic Tone", f"{a.get('academic_tone_score','N/A')}/10")
            with c2:
                st.metric("Clarity", f"{a.get('clarity_score','N/A')}/10")

            if a.get("suggestions"):
                st.markdown("**✏️ Improvement Suggestions**")
                for s in a["suggestions"]:
                    st.markdown(f"• {s}")

            if a.get("improved_sentence"):
                st.markdown("**✨ Rewritten Example**")
                st.info(a["improved_sentence"])

            if a.get("grammar_notes"):
                st.markdown("**📋 Grammar Notes**")
                for note in a["grammar_notes"]:
                    st.markdown(f"• {note}")

# ═══════════════════════════════════════════════════════════════════════════════
# PAGE: TRANSLATION
# ═══════════════════════════════════════════════════════════════════════════════
elif page == "Translation":
    st.markdown('<div class="section-header">🌍 Translation</div>', unsafe_allow_html=True)
    st.markdown('<div class="tip-box">Translate your draft written in regional Indian languages into academic English.</div>', unsafe_allow_html=True)

    lang_map = {
        "Hindi": "hi",
        "Marathi": "mr",
        "Gujarati": "gu",
        "Tamil": "ta",
        "Bengali": "bn",
        "Telugu": "te",
        "Kannada": "kn",
    }
    source_lang = st.selectbox("Source Language", list(lang_map.keys()))
    source_text = st.text_area("Your text in regional language:", height=150, placeholder=f"Type or paste your text in {source_lang}…")

    if st.button("🌍 Translate to Academic English", type="primary", disabled=not source_text.strip()):
        if TRANSLATOR_OK:
            try:
                with st.spinner("Translating…"):
                    t = GoogleTranslator(source=lang_map[source_lang], target="en")
                    translated = t.translate(source_text)

                prompt = f"""Refine this text into formal academic English suitable for a research paper.
Do not add information not present in the original.

Text: {translated}

Return ONLY the refined academic text, no extra commentary."""
                with st.spinner("Refining for academic tone…"):
                    refined = call_gemini(prompt)
                col1, col2 = st.columns(2)
                with col1:
                    st.markdown("**Translated (raw)**")
                    st.text_area("", value=translated, height=150)
                with col2:
                    st.markdown("**Academic English (refined)**")
                    st.text_area("", value=refined, height=150, key="refined_output")
                st.download_button("📥 Download", refined, "translated.txt")
            except Exception as e:
                st.error(f"Translation error: {e}")
        else:
            # Fallback via Gemini
            with st.spinner("Translating via AI…"):
                prompt = f"""Translate the following {source_lang} text to formal academic English.
Do not add any information not present in the original.

Text: {source_text}

Return only the translated academic English text."""
                result = call_gemini(prompt)
            st.text_area("Academic English Translation", value=result, height=200)

# ═══════════════════════════════════════════════════════════════════════════════
# PAGE: EXPORT
# ═══════════════════════════════════════════════════════════════════════════════
elif page == "Export":
    st.markdown('<div class="section-header">📤 Export Your Paper</div>', unsafe_allow_html=True)

    topic = st.session_state.selected_topic
    sections = st.session_state.paper_sections
    refs = st.session_state.references

    if not sections:
        st.info("No sections generated yet. Use the Paper Builder to write your paper first.")
    else:
        section_order = ["Abstract", "Introduction", "Literature Review", "Methodology", "Results Framework", "Discussion", "Conclusion", "Future Work"]
        available = [s for s in section_order if s in sections]
        st.markdown(f"**Sections ready:** {', '.join(available)}")
        st.markdown(f"**References:** {len(refs)}")

        title = topic["title"] if topic else "My Research Paper"
        paper_title = st.text_input("Paper Title", value=title)
        author_name = st.text_input("Author Name(s)", placeholder="e.g. Ananya Sharma, Priya Mehta")
        institution = st.text_input("Institution", placeholder="e.g. Mumbai University, BCA Department")

        col1, col2 = st.columns(2)

        with col1:
            st.markdown("#### 📄 Export as DOCX")
            if DOCX_OK:
                if st.button("📄 Generate DOCX", type="primary"):
                    doc = DocxDocument()
                    doc.add_heading(paper_title, level=0)
                    if author_name:
                        doc.add_paragraph(author_name).alignment = 1
                    if institution:
                        doc.add_paragraph(institution).alignment = 1
                    doc.add_paragraph(f"Generated: {datetime.now().strftime('%B %Y')}").alignment = 1
                    doc.add_page_break()

                    for sec in section_order:
                        if sec in sections:
                            doc.add_heading(sec, level=1)
                            doc.add_paragraph(sections[sec])

                    if refs:
                        doc.add_heading("References", level=1)
                        for r in refs:
                            authors_str = ", ".join(r["authors"][:3])
                            doc.add_paragraph(
                                f"{authors_str} ({r['year']}). {r['title']}. {r['source']}. {r['doi']}",
                                style="List Number",
                            )

                    buf = io.BytesIO()
                    doc.save(buf)
                    buf.seek(0)
                    st.download_button("⬇️ Download DOCX", buf, f"{paper_title[:40]}.docx",
                                       "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            else:
                st.warning("python-docx not installed.")

        with col2:
            st.markdown("#### 📋 Export as PDF")
            if REPORTLAB_OK:
                if st.button("📋 Generate PDF", type="primary"):
                    buf = io.BytesIO()
                    doc = SimpleDocTemplate(buf, pagesize=letter)
                    styles = getSampleStyleSheet()
                    story = []

                    story.append(Paragraph(paper_title, styles["Title"]))
                    if author_name:
                        story.append(Paragraph(author_name, styles["Normal"]))
                    if institution:
                        story.append(Paragraph(institution, styles["Normal"]))
                    story.append(Spacer(1, 20))

                    for sec in section_order:
                        if sec in sections:
                            story.append(Paragraph(sec, styles["Heading1"]))
                            story.append(Spacer(1, 6))
                            for para in sections[sec].split("\n\n"):
                                if para.strip():
                                    story.append(Paragraph(para.strip(), styles["BodyText"]))
                                    story.append(Spacer(1, 6))

                    if refs:
                        story.append(Paragraph("References", styles["Heading1"]))
                        for i, r in enumerate(refs, 1):
                            authors_str = ", ".join(r["authors"][:3])
                            story.append(Paragraph(
                                f"[{i}] {authors_str} ({r['year']}). {r['title']}. {r['source']}.",
                                styles["BodyText"]
                            ))

                    doc.build(story)
                    buf.seek(0)
                    st.download_button("⬇️ Download PDF", buf, f"{paper_title[:40]}.pdf", "application/pdf")
            else:
                st.warning("reportlab not installed.")

        st.divider()
        st.markdown("#### 📋 Preview")
        for sec in section_order:
            if sec in sections:
                with st.expander(sec):
                    st.write(sections[sec])

# ═══════════════════════════════════════════════════════════════════════════════
# footer
# ═══════════════════════════════════════════════════════════════════════════════
st.markdown("""
<div style="text-align:center;padding:2rem 0 1rem;color:#94A3B8;font-size:0.78rem">
    ResearchPaper AI Studio · Powered by Gemini 2.0 Flash · For students, by design<br>
    <span style="font-size:0.72rem">All citations come from real database searches. No information is invented.</span>
</div>
""", unsafe_allow_html=True)
